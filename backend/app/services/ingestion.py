import io
import re
from dataclasses import dataclass
from pathlib import Path

import fitz
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from ebooklib import ITEM_DOCUMENT
from ebooklib import epub as epub_reader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".epub", ".html", ".htm", ".azw"}
TEXT_RENDERABLE_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".epub", ".html", ".htm"}


@dataclass
class ExtractedSection:
    label: str
    content: str


@dataclass
class ExtractedDocument:
    extension: str
    text: str
    sections: list[ExtractedSection]


def validate_filename(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type: {ext}. Allowed: {allowed}")

    if ext == ".azw":
        raise ValueError(
            "AZW ingestion is not supported directly. "
            "Please convert Kindle AZW files to EPUB, PDF, or HTML first."
        )
    return ext


def extract_text_from_bytes(filename: str, data: bytes) -> ExtractedDocument:
    ext = validate_filename(filename)

    if ext == ".txt":
        text = data.decode("utf-8", errors="ignore")
        cleaned = clean_text(text)
        return ExtractedDocument(extension=ext, text=cleaned, sections=[ExtractedSection(label="Document", content=cleaned)])

    if ext == ".md":
        text = data.decode("utf-8", errors="ignore")
        sections = _split_markdown_sections(text)
        return ExtractedDocument(extension=ext, text=_sections_to_text(sections), sections=sections)

    if ext in {".html", ".htm"}:
        html = data.decode("utf-8", errors="ignore")
        sections = _extract_sections_from_html(html)
        return ExtractedDocument(extension=ext, text=_sections_to_text(sections), sections=sections)

    if ext == ".pdf":
        sections = _extract_sections_from_pdf(data)
        return ExtractedDocument(extension=ext, text=_sections_to_text(sections), sections=sections)

    if ext == ".docx":
        sections = _extract_sections_from_docx(data)
        return ExtractedDocument(extension=ext, text=_sections_to_text(sections), sections=sections)

    if ext == ".epub":
        sections = _extract_sections_from_epub(data)
        return ExtractedDocument(extension=ext, text=_sections_to_text(sections), sections=sections)

    raise ValueError("Could not extract text")


def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 200) -> list[str]:
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = max(0, end - overlap)
    return chunks


def _sections_to_text(sections: list[ExtractedSection]) -> str:
    if not sections:
        return ""
    return "\n\n".join(f"## {section.label}\n{section.content}" for section in sections)


def _split_markdown_sections(markdown_text: str) -> list[ExtractedSection]:
    sections: list[ExtractedSection] = []
    current_label = "Introduction"
    current_lines: list[str] = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if re.match(r"^\s*#{1,6}\s+", line):
            content = clean_text("\n".join(current_lines))
            if content:
                sections.append(ExtractedSection(label=current_label, content=content))
            current_label = re.sub(r"^\s*#{1,6}\s+", "", line).strip() or "Untitled Section"
            current_lines = []
            continue
        current_lines.append(line)

    trailing = clean_text("\n".join(current_lines))
    if trailing:
        sections.append(ExtractedSection(label=current_label, content=trailing))

    if not sections:
        cleaned = clean_text(markdown_text)
        if cleaned:
            sections.append(ExtractedSection(label="Document", content=cleaned))
    return sections


def _extract_sections_from_pdf(data: bytes) -> list[ExtractedSection]:
    doc = fitz.open(stream=data, filetype="pdf")
    sections: list[ExtractedSection] = []
    try:
        for page_index, page in enumerate(doc, start=1):
            text = clean_text(page.get_text("text"))
            if text:
                sections.append(ExtractedSection(label=f"Page {page_index}", content=text))
    finally:
        doc.close()

    if not sections:
        raise ValueError("Could not extract readable text from PDF.")
    return sections


def _extract_sections_from_docx(data: bytes) -> list[ExtractedSection]:
    document = DocxDocument(io.BytesIO(data))
    sections: list[ExtractedSection] = []

    current_label = "Document"
    current_lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            content = clean_text("\n".join(current_lines))
            if content:
                sections.append(ExtractedSection(label=current_label, content=content))
            current_label = text
            current_lines = []
            continue
        current_lines.append(text)

    trailing = clean_text("\n".join(current_lines))
    if trailing:
        sections.append(ExtractedSection(label=current_label, content=trailing))

    if not sections:
        raise ValueError("Could not extract readable text from DOCX.")
    return sections


def _extract_sections_from_html(html_text: str) -> list[ExtractedSection]:
    soup = BeautifulSoup(html_text, "html.parser")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()

    heading_tags = soup.find_all(re.compile(r"^h[1-6]$"))
    sections: list[ExtractedSection] = []

    if not heading_tags:
        cleaned = clean_text(soup.get_text("\n"))
        if cleaned:
            return [ExtractedSection(label="Document", content=cleaned)]
        return []

    for index, heading in enumerate(heading_tags, start=1):
        label = heading.get_text(" ", strip=True) or f"Section {index}"
        content_parts: list[str] = []
        for sibling in heading.next_siblings:
            sibling_name = getattr(sibling, "name", None)
            if sibling_name and re.match(r"^h[1-6]$", sibling_name):
                break
            if hasattr(sibling, "get_text"):
                text = sibling.get_text(" ", strip=True)
                if text:
                    content_parts.append(text)
            else:
                raw = str(sibling).strip()
                if raw:
                    content_parts.append(raw)

        content = clean_text("\n".join(content_parts))
        if content:
            sections.append(ExtractedSection(label=label, content=content))

    return sections


def _extract_sections_from_epub(data: bytes) -> list[ExtractedSection]:
    book = epub_reader.read_epub(io.BytesIO(data))
    sections: list[ExtractedSection] = []

    for index, item in enumerate(book.get_items_of_type(ITEM_DOCUMENT), start=1):
        html = item.get_content().decode("utf-8", errors="ignore")
        html_sections = _extract_sections_from_html(html)

        if not html_sections:
            continue

        # Keep chapter-level grouping by prefixing extracted labels.
        chapter_label = _guess_epub_chapter_label(item, index)
        for section in html_sections:
            merged_label = f"{chapter_label} · {section.label}" if section.label != "Document" else chapter_label
            sections.append(ExtractedSection(label=merged_label, content=section.content))

    if not sections:
        raise ValueError("Could not extract readable text from EPUB.")
    return sections


def _guess_epub_chapter_label(item: object, index: int) -> str:
    try:
        name = getattr(item, "file_name", None) or getattr(item, "get_name", lambda: None)()
    except Exception:
        name = None

    if name:
        base = Path(str(name)).stem.replace("_", " ").replace("-", " ").strip()
        if base:
            return base.title()
    return f"Chapter {index}"
